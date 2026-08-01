import pypdf
import docx
from deep_translator import GoogleTranslator

# Script de traduction automatisée du Manuel des Indicateurs ODD (362 pages)
# Découpage du texte par blocs pour éviter le dépassement de quota API

def translate_pdf_to_word(pdf_path, output_docx_path):
    reader = pypdf.PdfReader(pdf_path)
    doc = docx.Document()
    translator = GoogleTranslator(source='en', target='fr')
    
    print(f"Début de la traduction : {len(reader.pages)} pages à traiter...")
    
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if not text.strip():
            continue
            
        doc.add_heading(f"Page {page_num + 1}", level=2)
        
        # Découpage par paragraphes pour une meilleure stabilité de traduction
        paragraphs = text.split('\n\n')
        for p in paragraphs:
            if p.strip():
                try:
                    # Gestion de la limite de caractères de traduction (max 4500 caractères)
                    chunks = [p[i:i+4000] for i in range(0, len(p), 4000)]
                    translated_p = " ".join([translator.translate(chunk) for chunk in chunks])
                    doc.add_paragraph(translated_p)
                except Exception as e:
                    print(f"Erreur d'envoi page {page_num + 1}: {e}")
                    doc.add_paragraph(p) # Sauvegarde du texte brut en cas d'erreur de réseau
                    
        if (page_num + 1) % 10 == 0:
            print(f"Progression : {page_num + 1}/{len(reader.pages)} pages traitées...")
            
    doc.save(output_docx_path)
    print(f"Traduction terminée avec succès ! Fichier sauvegardé sous : {output_docx_path}")

if __name__ == "__main__":
    translate_pdf_to_word("Manuel des indicateurs ODD_Nations Unies_2026.pdf", "Manuel_Indicateurs_ODD_Traduit_FR.docx")
